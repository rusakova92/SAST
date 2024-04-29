import git
import subprocess
from docx import Document

# Клонируем репозиторий
repository_url = 'https://github.com/rusakova92/SAST.git'
repository_dir = 'repository_folder'
repo = git.Repo.clone_from(repository_url, repository_dir)

# Запуск SAST-анализатора
output = subprocess.check_output(['SAST_analyzer', repository_dir])

# Создание документ с результатами анализа
doc = Document()
doc.add_heading('Результаты анализа уязвимостей', level=1)
doc.add_paragraph('Репозиторий: {}'.format(repository_url))
doc.add_paragraph('Результаты анализа:')
doc.add_paragraph(output.decode('utf-8'))

doc.save('report.docx')
